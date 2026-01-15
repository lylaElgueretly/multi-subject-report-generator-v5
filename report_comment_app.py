# =========================================
# MULTI-SUBJECT REPORT COMMENT GENERATOR - Secure Streamlit Version
# Supports Year 5, 7 & 8; Subjects: English, Maths, Science
# =========================================

import random
import streamlit as st
import tempfile
import os
import time
from datetime import datetime, timedelta
import pandas as pd
import io
import re

# ========== DOCX IMPORT WITH FALLBACK ==========
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    Document = None

# ========== SECURITY & PRIVACY SETTINGS ==========
TARGET_CHARS = 499
MAX_FILE_SIZE_MB = 5
MAX_ROWS_PER_UPLOAD = 100

# ========== PAGE CONFIGURATION ==========
st.set_page_config(
    page_title="🔒 Secure Report Generator",
    page_icon="📚",
    layout="wide",
    initial_sidebar_state="expanded"
)

# ========== SECURITY INITIALIZATION ==========
if 'app_initialized' not in st.session_state:
    st.session_state.clear()
    st.session_state.app_initialized = True
    st.session_state.upload_count = 0
    st.session_state.last_upload_time = datetime.now()
    st.session_state.generated_files = []

# ========== IMPORT STATEMENTS FROM VARIANT FILES ==========
try:
    # Year 5 English - Use variant1 as default
    from statements_year5_English_variant1 import (
        opening_phrases as opening_5_eng,
        attitude_bank as attitude_5_eng,
        reading_bank as reading_5_eng,
        writing_bank as writing_5_eng,
        reading_target_bank as target_5_eng,
        writing_target_bank as target_write_5_eng,
        closer_bank as closer_5_eng
    )

    # Year 5 English variant2
    from statements_year5_English_variant2 import (
        opening_phrases as opening_5_eng_v2,
        attitude_bank as attitude_5_eng_v2,
        reading_bank as reading_5_eng_v2,
        writing_bank as writing_5_eng_v2,
        reading_target_bank as target_5_eng_v2,
        writing_target_bank as target_write_5_eng_v2,
        closer_bank as closer_5_eng_v2
    )

    # Year 5 Maths - Use variant1 as default
    from statements_year5_Maths_variant1 import (
        opening_phrases as opening_5_maths,
        attitude_bank as attitude_5_maths,
        number_bank as number_5_maths,
        problem_solving_bank as problem_5_maths,
        target_bank as target_5_maths,
        closer_bank as closer_5_maths
    )

    # Year 5 Maths variant2
    from statements_year5_Maths_variant2 import (
        opening_phrases as opening_5_maths_v2,
        attitude_bank as attitude_5_maths_v2,
        number_bank as number_5_maths_v2,
        problem_solving_bank as problem_5_maths_v2,
        target_bank as target_5_maths_v2,
        closer_bank as closer_5_maths_v2
    )

    # Year 5 Science - Use variant1 as default
    from statements_year5_Science_variant1 import (
        opening_phrases as opening_5_sci,
        attitude_bank as attitude_5_sci,
        science_bank as science_5_sci,
        target_bank as target_5_sci,
        closer_bank as closer_5_sci
    )

    # Year 5 Science variant2
    from statements_year5_Science_variant2 import (
        opening_phrases as opening_5_sci_v2,
        attitude_bank as attitude_5_sci_v2,
        science_bank as science_5_sci_v2,
        target_bank as target_5_sci_v2,
        closer_bank as closer_5_sci_v2
    )

    # Year 7 English - Use variant1 as default
    from statements_year7_English_variant1 import (
        opening_phrases as opening_7_eng,
        attitude_bank as attitude_7_eng,
        reading_bank as reading_7_eng,
        writing_bank as writing_7_eng,
        reading_target_bank as target_7_eng,
        writing_target_bank as target_write_7_eng,
        closer_bank as closer_7_eng
    )

    # Year 7 English variant2
    from statements_year7_English_variant2 import (
        opening_phrases as opening_7_eng_v2,
        attitude_bank as attitude_7_eng_v2,
        reading_bank as reading_7_eng_v2,
        writing_bank as writing_7_eng_v2,
        reading_target_bank as target_7_eng_v2,
        writing_target_bank as target_write_7_eng_v2,
        closer_bank as closer_7_eng_v2
    )

    # Year 7 Maths - Use variant1 as default
    from statements_year7_Maths_variant1 import (
        opening_phrases as opening_7_maths,
        attitude_bank as attitude_7_maths,
        number_and_algebra_bank as number_7_maths,
        geometry_and_measurement_bank as geometry_7_maths,
        problem_solving_and_reasoning_bank as problem_7_maths,
        target_bank as target_7_maths,
        closer_bank as closer_7_maths
    )

    # Year 7 Maths variant2
    from statements_year7_Maths_variant2 import (
        opening_phrases as opening_7_maths_v2,
        attitude_bank as attitude_7_maths_v2,
        number_and_algebra_bank as number_7_maths_v2,
        geometry_and_measurement_bank as geometry_7_maths_v2,
        problem_solving_and_reasoning_bank as problem_7_maths_v2,
        target_bank as target_7_maths_v2,
        closer_bank as closer_7_maths_v2
    )

    # Year 7 Science - Use variant1 as default
    from statements_year7_science_variant1 import (
        opening_phrases as opening_7_sci,
        attitude_bank as attitude_7_sci,
        science_bank as science_7_sci,
        target_bank as target_7_sci,
        closer_bank as closer_7_sci
    )

    # Year 7 Science variant2
    from statements_year7_science_variant2 import (
        opening_phrases as opening_7_sci_v2,
        attitude_bank as attitude_7_sci_v2,
        science_bank as science_7_sci_v2,
        target_bank as target_7_sci_v2,
        closer_bank as closer_7_sci_v2
    )

    # Year 8 English - Use variant1 as default
    from statements_year8_English_variant1 import (
        opening_phrases as opening_8_eng,
        attitude_bank as attitude_8_eng,
        reading_bank as reading_8_eng,
        writing_bank as writing_8_eng,
        reading_target_bank as target_8_eng,
        writing_target_bank as target_write_8_eng,
        closer_bank as closer_8_eng
    )

    # Year 8 English variant2
    from statements_year8_English_variant2 import (
        opening_phrases as opening_8_eng_v2,
        attitude_bank as attitude_8_eng_v2,
        reading_bank as reading_8_eng_v2,
        writing_bank as writing_8_eng_v2,
        reading_target_bank as target_8_eng_v2,
        writing_target_bank as target_write_8_eng_v2,
        closer_bank as closer_8_eng_v2
    )

    # Year 8 Maths - Use variant1 as default
    from statements_year8_Maths_variant1 import (
        opening_phrases as opening_8_maths,
        attitude_bank as attitude_8_maths,
        maths_bank as maths_8_maths,
        target_bank as target_8_maths,
        closer_bank as closer_8_maths
    )

    # Year 8 Maths variant2
    from statements_year8_Maths_variant2 import (
        opening_phrases as opening_8_maths_v2,
        attitude_bank as attitude_8_maths_v2,
        maths_bank as maths_8_maths_v2,
        target_bank as target_8_maths_v2,
        closer_bank as closer_8_maths_v2
    )

    # Year 8 Science - Use variant1 as default
    from statements_year8_science_variant1 import (
        opening_phrases as opening_8_sci,
        attitude_bank as attitude_8_sci,
        science_bank as science_8_sci,
        target_bank as target_8_sci,
        closer_bank as closer_8_sci
    )

    # Year 8 Science variant2
    from statements_year8_science_variant2 import (
        opening_phrases as opening_8_sci_v2,
        attitude_bank as attitude_8_sci_v2,
        science_bank as science_8_sci_v2,
        target_bank as target_8_sci_v2,
        closer_bank as closer_8_sci_v2
    )

except ImportError as e:
    st.error(f"Missing required statement files: {e}")
    st.stop()

# ========== SECURITY FUNCTIONS ==========
def sanitize_input(text, max_length=100):
    """Sanitize user input to prevent injection attacks"""
    if not text:
        return ""
    sanitized = ''.join(c for c in text if c.isalnum() or c in " .'-")
    return sanitized[:max_length].strip().title()

def validate_file(file):
    """Validate uploaded file size and type"""
    if file.size > MAX_FILE_SIZE_MB * 1024 * 1024:
        return False, f"File too large (max {MAX_FILE_SIZE_MB}MB)"
    if not file.name.lower().endswith('.csv'):
        return False, "Only CSV files allowed"
    return True, ""

def process_csv_securely(uploaded_file):
    """Process CSV with auto-cleanup of temp files"""
    with tempfile.NamedTemporaryFile(delete=False, suffix='.csv', mode='wb') as tmp:
        tmp.write(uploaded_file.getvalue())
        temp_path = tmp.name

    try:
        df = pd.read_csv(temp_path, nrows=MAX_ROWS_PER_UPLOAD + 1)
        if len(df) > MAX_ROWS_PER_UPLOAD:
            st.warning(f"Only processing first {MAX_ROWS_PER_UPLOAD} rows")
            df = df.head(MAX_ROWS_PER_UPLOAD)
        if 'Student Name' in df.columns:
            df['Student Name'] = df['Student Name'].apply(lambda x: sanitize_input(str(x)))
        return df
    except Exception as e:
        st.error(f"Error reading CSV: {e}")
        return None
    finally:
        try:
            os.unlink(temp_path)
        except:
            pass

# ========== HELPER FUNCTIONS ==========
def get_pronouns(gender):
    gender = gender.lower()
    if gender == "male":
        return "he", "his"
    elif gender == "female":
        return "she", "her"
    return "they", "their"

def lowercase_first(text):
    return text[0].lower() + text[1:] if text else ""

def truncate_comment(comment, target=TARGET_CHARS):
    if len(comment) <= target:
        return comment
    truncated = comment[:target].rstrip(" ,;.")
    if "." in truncated:
        truncated = truncated[:truncated.rfind(".")+1]
    return truncated

def fix_pronouns_in_text(text, pronoun, possessive):
    """Fix gender pronouns in statement text using word boundaries"""
    if not text:
        return text

    text = re.sub(r'\bhe\b', pronoun, text, flags=re.IGNORECASE)
    text = re.sub(r'\bHe\b', pronoun.capitalize(), text)
    text = re.sub(r'\bhis\b', possessive, text, flags=re.IGNORECASE)
    text = re.sub(r'\bHis\b', possessive.capitalize(), text)
    text = re.sub(r'\bhim\b', pronoun, text, flags=re.IGNORECASE)
    text = re.sub(r'\bHim\b', pronoun.capitalize(), text)
    text = re.sub(r'\bhimself\b', f"{pronoun}self", text, flags=re.IGNORECASE)
    text = re.sub(r'\bherself\b', f"{pronoun}self", text, flags=re.IGNORECASE)

    return text

# ========== COMMENT GENERATOR ==========
def generate_comment(subject, year, name, gender, att, achieve, target, pronouns, attitude_target=None, variant=1):
    """
    Generate a report comment.
    
    Args:
        variant (int): 1 = variant1 (default), 2 = variant2
    """
    p, p_poss = pronouns
    name = sanitize_input(name)

    # Select statements based on variant
    if variant == 1:
        # Use variant1 (default)
        if year == 5:
            if subject == "English":
                opening = random.choice(opening_5_eng)
                attitude_text = fix_pronouns_in_text(attitude_5_eng[att], p, p_poss)
                reading_text = fix_pronouns_in_text(reading_5_eng[achieve], p, p_poss)
                writing_text = fix_pronouns_in_text(writing_5_eng[achieve], p, p_poss)
                reading_target_text = fix_pronouns_in_text(target_5_eng[target], p, p_poss)
                writing_target_text = fix_pronouns_in_text(target_write_5_eng[target], p, p_poss)
                closer_sentence = random.choice(closer_5_eng)
            elif subject == "Maths":
                opening = random.choice(opening_5_maths)
                attitude_text = fix_pronouns_in_text(attitude_5_maths[att], p, p_poss)
                reading_text = fix_pronouns_in_text(number_5_maths[achieve], p, p_poss)
                reading_target_text = fix_pronouns_in_text(target_5_maths[target], p, p_poss)
                writing_text = ""
                writing_target_text = ""
                closer_sentence = random.choice(closer_5_maths)
            else:  # Science
                opening = random.choice(opening_5_sci)
                attitude_text = fix_pronouns_in_text(attitude_5_sci[att], p, p_poss)
                reading_text = fix_pronouns_in_text(science_5_sci[achieve], p, p_poss)
                reading_target_text = fix_pronouns_in_text(target_5_sci[target], p, p_poss)
                writing_text = ""
                writing_target_text = ""
                closer_sentence = random.choice(closer_5_sci)
        
        elif year == 7:
            if subject == "English":
                opening = random.choice(opening_7_eng)
                attitude_text = fix_pronouns_in_text(attitude_7_eng[att], p, p_poss)
                reading_text = fix_pronouns_in_text(reading_7_eng[achieve], p, p_poss)
                writing_text = fix_pronouns_in_text(writing_7_eng[achieve], p, p_poss)
                reading_target_text = fix_pronouns_in_text(target_7_eng[target], p, p_poss)
                writing_target_text = fix_pronouns_in_text(target_write_7_eng[target], p, p_poss)
                closer_sentence = random.choice(closer_7_eng)
            elif subject == "Maths":
                opening = random.choice(opening_7_maths)
                attitude_text = fix_pronouns_in_text(attitude_7_maths[att], p, p_poss)
                reading_text = fix_pronouns_in_text(number_7_maths[achieve], p, p_poss)
                reading_target_text = fix_pronouns_in_text(target_7_maths[target], p, p_poss)
                writing_text = ""
                writing_target_text = ""
                closer_sentence = random.choice(closer_7_maths)
            else:  # Science
                opening = random.choice(opening_7_sci)
                attitude_text = fix_pronouns_in_text(attitude_7_sci[att], p, p_poss)
                reading_text = fix_pronouns_in_text(science_7_sci[achieve], p, p_poss)
                reading_target_text = fix_pronouns_in_text(target_7_sci[target], p, p_poss)
                writing_text = ""
                writing_target_text = ""
                closer_sentence = random.choice(closer_7_sci)
        
        else:  # year == 8
            if subject == "English":
                opening = random.choice(opening_8_eng)
                attitude_text = fix_pronouns_in_text(attitude_8_eng[att], p, p_poss)
                reading_text = fix_pronouns_in_text(reading_8_eng[achieve], p, p_poss)
                writing_text = fix_pronouns_in_text(writing_8_eng[achieve], p, p_poss)
                reading_target_text = fix_pronouns_in_text(target_8_eng[target], p, p_poss)
                writing_target_text = fix_pronouns_in_text(target_write_8_eng[target], p, p_poss)
                closer_sentence = random.choice(closer_8_eng)
            elif subject == "Maths":
                opening = random.choice(opening_8_maths)
                attitude_text = fix_pronouns_in_text(attitude_8_maths[att], p, p_poss)
                reading_text = fix_pronouns_in_text(maths_8_maths[achieve], p, p_poss)
                reading_target_text = fix_pronouns_in_text(target_8_maths[target], p, p_poss)
                writing_text = ""
                writing_target_text = ""
                closer_sentence = random.choice(closer_8_maths)
            else:  # Science
                opening = random.choice(opening_8_sci)
                attitude_text = fix_pronouns_in_text(attitude_8_sci[att], p, p_poss)
                reading_text = fix_pronouns_in_text(science_8_sci[achieve], p, p_poss)
                reading_target_text = fix_pronouns_in_text(target_8_sci[target], p, p_poss)
                writing_text = ""
                writing_target_text = ""
                closer_sentence = random.choice(closer_8_sci)
    
    else:  # variant == 2
        # Use variant2
        if year == 5:
            if subject == "English":
                opening = random.choice(opening_5_eng_v2)
                attitude_text = fix_pronouns_in_text(attitude_5_eng_v2[att], p, p_poss)
                reading_text = fix_pronouns_in_text(reading_5_eng_v2[achieve], p, p_poss)
                writing_text = fix_pronouns_in_text(writing_5_eng_v2[achieve], p, p_poss)
                reading_target_text = fix_pronouns_in_text(target_5_eng_v2[target], p, p_poss)
                writing_target_text = fix_pronouns_in_text(target_write_5_eng_v2[target], p, p_poss)
                closer_sentence = random.choice(closer_5_eng_v2)
            elif subject == "Maths":
                opening = random.choice(opening_5_maths_v2)
                attitude_text = fix_pronouns_in_text(attitude_5_maths_v2[att], p, p_poss)
                reading_text = fix_pronouns_in_text(number_5_maths_v2[achieve], p, p_poss)
                reading_target_text = fix_pronouns_in_text(target_5_maths_v2[target], p, p_poss)
                writing_text = ""
                writing_target_text = ""
                closer_sentence = random.choice(closer_5_maths_v2)
            else:  # Science
                opening = random.choice(opening_5_sci_v2)
                attitude_text = fix_pronouns_in_text(attitude_5_sci_v2[att], p, p_poss)
                reading_text = fix_pronouns_in_text(science_5_sci_v2[achieve], p, p_poss)
                reading_target_text = fix_pronouns_in_text(target_5_sci_v2[target], p, p_poss)
                writing_text = ""
                writing_target_text = ""
                closer_sentence = random.choice(closer_5_sci_v2)
        
        elif year == 7:
            if subject == "English":
                opening = random.choice(opening_7_eng_v2)
                attitude_text = fix_pronouns_in_text(attitude_7_eng_v2[att], p, p_poss)
                reading_text = fix_pronouns_in_text(reading_7_eng_v2[achieve], p, p_poss)
                writing_text = fix_pronouns_in_text(writing_7_eng_v2[achieve], p, p_poss)
                reading_target_text = fix_pronouns_in_text(target_7_eng_v2[target], p, p_poss)
                writing_target_text = fix_pronouns_in_text(target_write_7_eng_v2[target], p, p_poss)
                closer_sentence = random.choice(closer_7_eng_v2)
            elif subject == "Maths":
                opening = random.choice(opening_7_maths_v2)
                attitude_text = fix_pronouns_in_text(attitude_7_maths_v2[att], p, p_poss)
                reading_text = fix_pronouns_in_text(number_7_maths_v2[achieve], p, p_poss)
                reading_target_text = fix_pronouns_in_text(target_7_maths_v2[target], p, p_poss)
                writing_text = ""
                writing_target_text = ""
                closer_sentence = random.choice(closer_7_maths_v2)
            else:  # Science
                opening = random.choice(opening_7_sci_v2)
                attitude_text = fix_pronouns_in_text(attitude_7_sci_v2[att], p, p_poss)
                reading_text = fix_pronouns_in_text(science_7_sci_v2[achieve], p, p_poss)
                reading_target_text = fix_pronouns_in_text(target_7_sci_v2[target], p, p_poss)
                writing_text = ""
                writing_target_text = ""
                closer_sentence = random.choice(closer_7_sci_v2)
        
        else:  # year == 8
            if subject == "English":
                opening = random.choice(opening_8_eng_v2)
                attitude_text = fix_pronouns_in_text(attitude_8_eng_v2[att], p, p_poss)
                reading_text = fix_pronouns_in_text(reading_8_eng_v2[achieve], p, p_poss)
                writing_text = fix_pronouns_in_text(writing_8_eng_v2[achieve], p, p_poss)
                reading_target_text = fix_pronouns_in_text(target_8_eng_v2[target], p, p_poss)
                writing_target_text = fix_pronouns_in_text(target_write_8_eng_v2[target], p, p_poss)
                closer_sentence = random.choice(closer_8_eng_v2)
            elif subject == "Maths":
                opening = random.choice(opening_8_maths_v2)
                attitude_text = fix_pronouns_in_text(attitude_8_maths_v2[att], p, p_poss)
                reading_text = fix_pronouns_in_text(maths_8_maths_v2[achieve], p, p_poss)
                reading_target_text = fix_pronouns_in_text(target_8_maths_v2[target], p, p_poss)
                writing_text = ""
                writing_target_text = ""
                closer_sentence = random.choice(closer_8_maths_v2)
            else:  # Science
                opening = random.choice(opening_8_sci_v2)
                attitude_text = fix_pronouns_in_text(attitude_8_sci_v2[att], p, p_poss)
                reading_text = fix_pronouns_in_text(science_8_sci_v2[achieve], p, p_poss)
                reading_target_text = fix_pronouns_in_text(target_8_sci_v2[target], p, p_poss)
                writing_text = ""
                writing_target_text = ""
                closer_sentence = random.choice(closer_8_sci_v2)

    # Construct sentences
    attitude_sentence = f"{opening} {name} {attitude_text}"
    if not attitude_sentence.endswith('.'):
        attitude_sentence += '.'

    if reading_text and reading_text[0].islower():
        reading_text = f"{p} {reading_text}"
    
    if subject == "English":
        reading_sentence = f"In reading, {reading_text}" if reading_text else ""
    else:
        reading_sentence = reading_text if reading_text else ""
    
    if not reading_sentence.endswith('.') and reading_sentence:
        reading_sentence += '.'

    if writing_text:
        if writing_text[0].islower():
            writing_text = f"{p} {writing_text}"
        writing_sentence = f"In writing, {writing_text}"
        if not writing_sentence.endswith('.'):
            writing_sentence += '.'
    else:
        writing_sentence = ""

    if reading_target_text:
        reading_target_sentence = f"For the next term, {p} should {lowercase_first(reading_target_text)}"
        if not reading_target_sentence.endswith('.'):
            reading_target_sentence += '.'
    else:
        reading_target_sentence = ""

    if writing_target_text:
        writing_target_sentence = f"Additionally, {p} should {lowercase_first(writing_target_text)}"
        if not writing_target_sentence.endswith('.'):
            writing_target_sentence += '.'
    else:
        writing_target_sentence = ""

    # Optional attitude target
    if attitude_target:
        attitude_target = sanitize_input(attitude_target)
        attitude_target_sentence = f"{lowercase_first(attitude_target)}"
        if not attitude_target_sentence.endswith('.'):
            attitude_target_sentence += '.'
        attitude_target_sentence = attitude_target_sentence.replace('..', '.')
    else:
        attitude_target_sentence = ""

    # Assemble comment
    comment_parts = [
        attitude_sentence,
        reading_sentence,
        writing_sentence,
        reading_target_sentence,
        writing_target_sentence,
        closer_sentence,
        attitude_target_sentence
    ]

    comment = " ".join([c for c in comment_parts if c])
    comment = comment.strip()

    # Fix punctuation
    if not comment.endswith('.'):
        comment += '.'
    comment = comment.replace('..', '.')

    comment = truncate_comment(comment, TARGET_CHARS)

    # Double-check ending punctuation after truncation
    if not comment.endswith('.'):
        comment = comment.rstrip(' ,;') + '.'
    comment = comment.replace('..', '.')

    return comment

# ========== STREAMLIT APP LAYOUT ==========

# Sidebar
with st.sidebar:
    st.title("📚 Navigation")

    app_mode = st.radio(
        "Choose Mode",
        ["Single Student", "Batch Upload", "Privacy Info"]
    )

    st.markdown("---")
    st.markdown("### 🔒 Privacy Features")
    st.info("""
    - No data stored on servers
    - All processing in memory
    - Auto-deletion of temp files
    - Input sanitization
    - Rate limiting enabled
    """)

    if st.button("🔄 Clear All Data", type="secondary", use_container_width=True):
        st.session_state.clear()
        st.session_state.app_initialized = True
        st.session_state.upload_count = 0
        st.session_state.last_upload_time = datetime.now()
        st.session_state.current_student = {}
        st.session_state.selected_comments = []
        st.success("All data cleared!")
        st.rerun()

    st.markdown("---")
    st.caption("v3.0 • Multi-Year Edition")

# Main content
col1, col2 = st.columns([1, 4])

with col1:
    try:
        st.image("logo.png", use_column_width=True)
    except:
        st.markdown("""
        <div style='text-align: center;'>
            <div style='font-size: 72px;'>📚</div>
        </div>
        """, unsafe_allow_html=True)

with col2:
    st.title("Multi-Subject Report Comment Generator")
    st.caption("~499 characters • Years 5, 7 & 8 • English, Maths, Science")

st.warning("""
**PRIVACY NOTICE:** All data is processed in memory only. No files are stored on our servers.
Close browser tab to completely erase all data.
""", icon="🔒")

# Progress tracker
st.subheader("🎯 Three Easy Steps")

if 'progress' not in st.session_state:
    st.session_state.progress = 1

step_col1, step_col2, step_col3 = st.columns(3)

def step_box(col, step_num, title, description):
    with col:
        is_current = st.session_state.progress == step_num
        bg_color = '#e6f3ff' if is_current else '#f8f9fa'
        st.markdown(f"""
        <div style='
            text-align: center;
            padding: 8px 5px;
            margin: 2px 0;
            background-color: {bg_color};
            border-radius: 8px;
            border-left: 4px solid #1E88E5;
            font-size: 0.9em;
        '>
            <div style='font-size: 1.2em; margin-bottom: 2px;'>
                {'✅' if st.session_state.progress > step_num else f'{step_num}.'} {title}
            </div>
            <div style='font-size: 0.85em; color: #666;'>{description}</div>
        </div>
        """, unsafe_allow_html=True)

step_box(step_col1, 1, "Select", "Choose student details")
step_box(step_col2, 2, "Generate & Approve", "Create and approve comments")
step_box(step_col3, 3, "Download", "Export your reports")

st.markdown("<br>", unsafe_allow_html=True)

# Initialize session state
if 'current_student' not in st.session_state:
    st.session_state.current_student = {}
if 'selected_comments' not in st.session_state:
    st.session_state.selected_comments = []

# ========== SINGLE STUDENT MODE ==========
if app_mode == "Single Student":
    st.subheader("👤 Single Student Entry")

    # Add reset settings option
    col_header1, col_header2 = st.columns([3, 1])
    with col_header2:
        if st.button("🔄 Reset Settings", help="Clear saved Subject/Year", use_container_width=True):
            st.session_state.last_subject = "English"
            st.session_state.last_year = 7
            st.success("Settings reset!")
            st.rerun()

    # Initialize persistent settings
    if 'last_subject' not in st.session_state:
        st.session_state.last_subject = "English"
    if 'last_year' not in st.session_state:
        st.session_state.last_year = 7

    with st.form("single_student_form", clear_on_submit=True):
        col1, col2 = st.columns(2)

        with col1:
            # Use last selected values as defaults
            subject_options = ["English", "Maths", "Science"]
            subject_index = subject_options.index(st.session_state.last_subject) if st.session_state.last_subject in subject_options else 0
            subject = st.selectbox("Subject", subject_options, index=subject_index)

            year_options = [5, 7, 8]
            year_index = year_options.index(st.session_state.last_year) if st.session_state.last_year in year_options else 1
            year = st.selectbox("Year", year_options, index=year_index)

            name = st.text_input("Student Name", placeholder="Enter first name only",
                                 key='student_name_input')
            gender = st.selectbox("Gender", ["Male", "Female"])

        with col2:
            att = st.selectbox("Attitude Band",
                             options=[90,85,80,75,70,65,60,55,40],
                             index=3)

            achieve = st.selectbox("Achievement Band",
                                 options=[90,85,80,75,70,65,60,55,40],
                                 index=3)

            target = st.selectbox("Target Band",
                                options=[90,85,80,75,70,65,60,55,40],
                                index=3)

            # Show if settings are remembered
            if st.session_state.last_subject != "English" or st.session_state.last_year != 7:
                st.caption(f"✓ Using saved: {st.session_state.last_subject} Year {st.session_state.last_year}")
            else:
                st.caption("💡 Subject & Year will be remembered for next student")

        attitude_target = st.text_area("Optional Attitude Next Steps",
                                     placeholder="E.g., continue to participate actively in class discussions...",
                                     height=60,
                                     key='attitude_target_input')

        col_submit = st.columns([4, 1])
        with col_submit[1]:
            submitted = st.form_submit_button("🚀 Generate Comment (Variant 1)", use_container_width=True)

    if submitted and name:
        # Save settings for next student
        st.session_state.last_subject = subject
        st.session_state.last_year = year

        name = sanitize_input(name)
        pronouns = get_pronouns(gender)

        with st.spinner("Generating comment..."):
            # Generate Variant 1 only
            comment_v1 = generate_comment(subject, year, name, gender, att, achieve,
                                         target, pronouns,
                                         st.session_state.get('attitude_target_input', ''),
                                         variant=1)
            
            # Store in session state
            st.session_state.current_student = {
                'name': name,
                'subject': subject,
                'year': year,
                'gender': gender,
                'att': att,
                'achieve': achieve,
                'target': target,
                'attitude_target': st.session_state.get('attitude_target_input', ''),
                'variant1': comment_v1,
                'variant2': None,  # Will be generated if requested
                'variant1_approved': False,
                'variant2_approved': False
            }

        st.session_state.progress = 2
        st.rerun()

    # Display generated Variant 1 if it exists
    if st.session_state.current_student and 'variant1' in st.session_state.current_student:
        current = st.session_state.current_student
        name = current['name']
        subject = current['subject']
        year = current['year']
        comment_v1 = current['variant1']
        
        st.subheader(f"📝 Generated Comment for {name} ({subject} Year {year})")
        
        # Display Variant 1
        st.markdown("### Variant 1 (Default)")
        st.text_area("Variant 1 Comment", comment_v1, height=150, key="variant1_display")
        
        # Character count
        char_count_v1 = len(comment_v1)
        st.caption(f"Characters: {char_count_v1}/{TARGET_CHARS}")
        
        # Action buttons for Variant 1
        col1_actions, col2_actions = st.columns([1, 1])
        
        with col1_actions:
            if not current['variant1_approved']:
                if st.button("✅ Approve Variant 1", type="primary", use_container_width=True):
                    # Add to selected comments
                    student_entry = {
                        'name': name,
                        'subject': subject,
                        'year': year,
                        'comment': comment_v1,
                        'variant': 'Variant 1',
                        'timestamp': datetime.now().strftime("%Y-%m-%d %H:%M")
                    }
                    st.session_state.selected_comments.append(student_entry)
                    st.session_state.current_student['variant1_approved'] = True
                    st.success(f"✓ Variant 1 approved for {name}!")
                    st.rerun()
            else:
                st.success("✓ Variant 1 already approved")
                
        with col2_actions:
            if st.button("🔄 Regenerate Variant 1", type="secondary", use_container_width=True):
                pronouns = get_pronouns(current['gender'])
                new_comment = generate_comment(
                    subject, year, name, current['gender'], 
                    current['att'], current['achieve'], current['target'],
                    pronouns, current['attitude_target'], variant=1
                )
                st.session_state.current_student['variant1'] = new_comment
                st.session_state.current_student['variant1_approved'] = False
                st.success("Variant 1 regenerated!")
                st.rerun()
        
        st.markdown("---")
        
        # Check if Variant 2 has been generated
        if current['variant2']:
            st.markdown("### Variant 2 (Alternative)")
            st.text_area("Variant 2 Comment", current['variant2'], height=150, key="variant2_display")
            
            # Character count
            char_count_v2 = len(current['variant2'])
            st.caption(f"Characters: {char_count_v2}/{TARGET_CHARS}")
            
            # Action buttons for Variant 2
            col1_v2, col2_v2 = st.columns([1, 1])
            
            with col1_v2:
                if not current['variant2_approved']:
                    if st.button("✅ Approve Variant 2", type="primary", use_container_width=True):
                        # Add to selected comments
                        student_entry = {
                            'name': name,
                            'subject': subject,
                            'year': year,
                            'comment': current['variant2'],
                            'variant': 'Variant 2',
                            'timestamp': datetime.now().strftime("%Y-%m-%d %H:%M")
                        }
                        st.session_state.selected_comments.append(student_entry)
                        st.session_state.current_student['variant2_approved'] = True
                        st.success(f"✓ Variant 2 approved for {name}!")
                        st.rerun()
                else:
                    st.success("✓ Variant 2 already approved")
                    
            with col2_v2:
                if st.button("🔄 Regenerate Variant 2", type="secondary", use_container_width=True):
                    pronouns = get_pronouns(current['gender'])
                    new_comment = generate_comment(
                        subject, year, name, current['gender'], 
                        current['att'], current['achieve'], current['target'],
                        pronouns, current['attitude_target'], variant=2
                    )
                    st.session_state.current_student['variant2'] = new_comment
                    st.session_state.current_student['variant2_approved'] = False
                    st.success("Variant 2 regenerated!")
                    st.rerun()
        
        # Generate Variant 2 button (only if not already generated)
        if not current['variant2']:
            st.markdown("---")
            if st.button("✨ Generate Variant 2 (Alternative)", type="secondary", use_container_width=True):
                pronouns = get_pronouns(current['gender'])
                with st.spinner("Generating Variant 2..."):
                    comment_v2 = generate_comment(
                        subject, year, name, current['gender'], 
                        current['att'], current['achieve'], current['target'],
                        pronouns, current['attitude_target'], variant=2
                    )
                    st.session_state.current_student['variant2'] = comment_v2
                    st.success("Variant 2 generated!")
                    st.rerun()
        
        # Navigation buttons
        st.markdown("---")
        col_nav1, col_nav2, col_nav3 = st.columns([1, 1, 1])
        
        with col_nav1:
            if st.button("➕ Add Another Student", type="primary", use_container_width=True):
                # Clear current student but keep selected comments
                st.session_state.current_student = {}
                if 'student_name_input' in st.session_state:
                    st.session_state.student_name_input = ""
                if 'attitude_target_input' in st.session_state:
                    st.session_state.attitude_target_input = ""
                st.session_state.progress = 1
                st.rerun()
                
        with col_nav2:
            if st.button("📋 Copy Variant 1", type="secondary", use_container_width=True):
                st.code(comment_v1, language=None)
                st.success("✓ Variant 1 copied to clipboard!")
                
        with col_nav3:
            if current['variant2']:
                if st.button("📋 Copy Variant 2", type="secondary", use_container_width=True):
                    st.code(current['variant2'], language=None)
                    st.success("✓ Variant 2 copied to clipboard!")

# ========== BATCH UPLOAD MODE ==========
elif app_mode == "Batch Upload":
    st.subheader("📁 Batch Upload (CSV)")

    st.info("""
    **CSV Format Required:**
    - Columns: Student Name, Gender, Subject, Year, Attitude, Achievement, Target
    - Gender: Male/Female
    - Subject: English/Maths/Science
    - Year: 5, 7, or 8
    - Bands: 90,85,80,75,70,65,60,55,40
    """)

    example_csv = """Student Name,Gender,Subject,Year,Attitude,Achievement,Target
Aseel,Female,English,5,75,80,85
Mohamed,Male,Maths,7,80,75,80
Sarah,Female,Science,8,85,90,85"""

    st.download_button(
        label="📥 Download Example CSV",
        data=example_csv,
        file_name="example_students.csv",
        mime="text/csv"
    )

    uploaded_file = st.file_uploader("Choose CSV file", type=['csv'])

    if uploaded_file:
        is_valid, msg = validate_file(uploaded_file)
        if not is_valid:
            st.error(msg)
            st.stop()

        with st.spinner("Processing CSV securely..."):
            df = process_csv_securely(uploaded_file)

        if df is not None:
            st.success(f"Processed {len(df)} students successfully")

            with st.expander("📋 Preview Data (First 5 rows)"):
                st.dataframe(df.head())

            if st.button("🚀 Generate All Comments (Variant 1)", type="primary"):
                if 'selected_comments' not in st.session_state:
                    st.session_state.selected_comments = []

                progress_bar = st.progress(0)
                status_text = st.empty()

                for idx, row in df.iterrows():
                    progress = (idx + 1) / len(df)
                    progress_bar.progress(progress)
                    status_text.text(f"Processing {idx + 1}/{len(df)}: {row.get('Student Name', 'Student')}")

                    try:
                        pronouns = get_pronouns(str(row.get('Gender', '')).lower())
                        comment = generate_comment(
                            subject=str(row.get('Subject', 'English')),
                            year=int(row.get('Year', 7)),
                            name=str(row.get('Student Name', '')),
                            gender=str(row.get('Gender', '')),
                            att=int(row.get('Attitude', 75)),
                            achieve=int(row.get('Achievement', 75)),
                            target=int(row.get('Target', 75)),
                            pronouns=pronouns,
                            variant=1
                        )

                        student_entry = {
                            'name': sanitize_input(str(row.get('Student Name', ''))),
                            'subject': str(row.get('Subject', 'English')),
                            'year': int(row.get('Year', 7)),
                            'comment': comment,
                            'variant': 'Variant 1',
                            'timestamp': datetime.now().strftime("%Y-%m-%d %H:%M")
                        }
                        st.session_state.selected_comments.append(student_entry)

                    except Exception as e:
                        st.error(f"Error processing row {idx + 1}: {e}")

                progress_bar.empty()
                status_text.empty()
                st.session_state.progress = 2
                st.success(f"Generated {len(df)} comments (Variant 1)!")
                st.session_state.last_upload_time = datetime.now()

# ========== PRIVACY INFO MODE ==========
elif app_mode == "Privacy Info":
    st.subheader("🔐 Privacy & Security Information")

    st.markdown("""
    ### How We Protect Student Data

    **Data Handling:**
    - All processing happens in your browser's memory
    - No student data is sent to or stored on our servers
    - Temporary files are created and immediately deleted
    - No database or persistent storage is used

    **Security Features:**
    1. **Input Sanitization** - Removes special characters from names
    2. **File Validation** - Checks file size and type
    3. **Auto-Cleanup** - Temporary files deleted after processing
    4. **Memory Clearing** - All data erased on browser close

    **Best Practices for Users:**
    - Use only first names or student IDs
    - Close browser tab when finished to clear all data
    - Download reports immediately after generation
    - For maximum privacy, use on school-managed devices

    **Compliance:**
    - Designed for use with anonymized data
    - Suitable for FERPA/GDPR compliant workflows
    - No third-party data sharing
    """)

    if st.button("🖨️ Print Privacy Notice", type="secondary"):
        privacy_text = """
        MULTI-SUBJECT REPORT GENERATOR - PRIVACY NOTICE

        Data Processing: All student data is processed locally in memory only.
        No data is transmitted to external servers or stored permanently.

        Data Retention: All data is cleared when the browser tab is closed.

        Security: Input sanitization and validation prevents data injection.

        Usage: For use with anonymized student data only.
        """
        st.text_area("Privacy Notice for Records", privacy_text, height=300)

# ========== DOWNLOAD SECTION ==========
if st.session_state.selected_comments:
    st.session_state.progress = 3
    st.markdown("---")
    st.subheader("📥 Download Reports")

    total_comments = len(st.session_state.selected_comments)
    st.info(f"You have {total_comments} approved comment(s) ready for download")

    with st.expander(f"👁️ Preview Approved Comments ({total_comments})"):
        for idx, entry in enumerate(st.session_state.selected_comments, 1):
            variant_label = f" ({entry.get('variant', '')})" if 'variant' in entry else ""
            st.markdown(f"**{idx}. {entry['name']}** ({entry['subject']} Year {entry['year']}{variant_label})")
            st.write(entry['comment'])
            st.caption(f"Added: {entry['timestamp']}")
            st.markdown("---")

    col_dl1, col_dl2, col_dl3 = st.columns(3)

    with col_dl1:
        if DOCX_AVAILABLE:
            if st.button("📄 Word Document", use_container_width=True):
                doc = Document()
                doc.add_heading('Report Comments', 0)
                doc.add_paragraph(f'Generated: {datetime.now().strftime("%Y-%m-%d %H:%M")}')
                doc.add_paragraph(f'Total Students: {total_comments}')
                doc.add_paragraph('')

                for entry in st.session_state.selected_comments:
                    variant_label = f" ({entry.get('variant', '')})" if 'variant' in entry else ""
                    doc.add_heading(f"{entry['name']} - {entry['subject']} Year {entry['year']}{variant_label}", level=2)
                    doc.add_paragraph(entry['comment'])
                    doc.add_paragraph('')

                bio = io.BytesIO()
                doc.save(bio)

                st.download_button(
                    label="⬇️ Download Word File",
                    data=bio.getvalue(),
                    file_name=f"report_comments_{datetime.now().strftime('%Y%m%d_%H%M')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )
        else:
            st.info("Word export requires python-docx package")

    with col_dl2:
        if st.button("📊 CSV Export", use_container_width=True):
            csv_data = []
            for entry in st.session_state.selected_comments:
                csv_data.append({
                    'Student Name': entry['name'],
                    'Subject': entry['subject'],
                    'Year': entry['year'],
                    'Variant': entry.get('variant', 'Variant 1'),
                    'Comment': entry['comment'],
                    'Generated': entry['timestamp']
                })

            df_export = pd.DataFrame(csv_data)
            csv_bytes = df_export.to_csv(index=False).encode('utf-8')

            st.download_button(
                label="⬇️ Download CSV",
                data=csv_bytes,
                file_name=f"report_comments_{datetime.now().strftime('%Y%m%d_%H%M')}.csv",
                mime="text/csv",
                use_container_width=True
            )

    with col_dl3:
        if st.button("🗑️ Clear & Start Over", type="secondary", use_container_width=True):
            st.session_state.selected_comments = []
            st.session_state.current_student = {}
            st.session_state.progress = 1
            st.success("All comments cleared! Ready for new entries.")
            if 'student_name_input' in st.session_state:
                st.session_state.student_name_input = ""
            if 'attitude_target_input' in st.session_state:
                st.session_state.attitude_target_input = ""
            st.rerun()

# ========== FOOTER ==========
st.markdown("---")
footer_cols = st.columns([2, 1])
with footer_cols[0]:
    st.caption("© Report Generator v3.0 • Multi-Year Edition")
with footer_cols[1]:
    if st.button("ℹ️ Quick Help", use_container_width=True):
        st.info("""
        **Quick Help:**
        1. **Select**: Choose student details
        2. **Generate**: Creates Variant 1 (default)
        3. **Approve**: Click "Approve Variant 1" to add to download list
        4. **Optional**: Generate Variant 2 if needed
        5. **Download**: Export approved comments

        **Features:**
        - Variant 1: Generated automatically
        - Variant 2: Optional alternative style
        - Approve individually: Choose which variants to keep
        - Regenerate: Get new versions if needed

        Need help? Contact support.
        """)
